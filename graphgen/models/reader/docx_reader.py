from typing import Any, Dict, List

from graphgen.bases.base_reader import BaseReader

try:
    from docx import Document
except ImportError:
    raise ImportError(
        "python-docx is required for reading Word documents. "
        "Please install it using: pip install python-docx"
    )


class DOCXReader(BaseReader):
    """
    Reader for Microsoft Word (.docx) files.
    Extracts text content from Word documents while preserving paragraph structure.
    Tables are converted to Markdown format for better structure preservation.
    """
    
    @staticmethod
    def _table_to_markdown(table) -> str:
        """
        Convert a Word table to Markdown format.
        
        :param table: docx table object
        :return: Markdown formatted table string
        """
        if len(table.rows) == 0:
            return ""
        
        # Extract all rows
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Get cell text, replace newlines with <br> or spaces, and escape pipe characters
                cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                cells.append(cell_text)
            # Only add non-empty rows
            if any(cell.strip() for cell in cells):
                rows.append(cells)
        
        if not rows:
            return ""
        
        # Determine number of columns (use max columns across all rows)
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols == 0:
            return ""
        
        # Normalize rows to have the same number of columns
        normalized_rows = []
        for row in rows:
            normalized_row = row + [""] * (max_cols - len(row))
            normalized_rows.append(normalized_row[:max_cols])  # Trim if too many columns
        
        # Build Markdown table
        markdown_lines = []
        
        # Header row (first row)
        if normalized_rows:
            header_row = normalized_rows[0]
            markdown_lines.append("| " + " | ".join(header_row) + " |")
            
            # Separator row
            markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")
            
            # Data rows (rest of rows)
            for row in normalized_rows[1:]:
                markdown_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(markdown_lines)
    
    def read(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Read text content from a .docx file.
        
        :param file_path: Path to the .docx file.
        :return: List containing a single dictionary with the document content.
        """
        try:
            doc = Document(file_path)
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraphs.append(para.text)
            
            # Join paragraphs with newlines to preserve structure
            full_text = "\n\n".join(paragraphs)
            
            # Also extract text from tables if present and convert to Markdown format
            tables_text = []
            for table in doc.tables:
                if len(table.rows) == 0:
                    continue
                
                try:
                    # Try to convert table to Markdown format
                    markdown_table = self._table_to_markdown(table)
                    if markdown_table:
                        tables_text.append(markdown_table)
                except Exception as e:
                    # If Markdown conversion fails, fallback to plain text format
                    try:
                        table_rows = []
                        for row in table.rows:
                            row_cells = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_cells.append(cell_text)
                            if row_cells:
                                table_rows.append(" | ".join(row_cells))
                        if table_rows:
                            tables_text.append("\n".join(table_rows))
                    except Exception:
                        # If even plain text extraction fails, skip this table
                        pass
            
            # Combine paragraphs and tables
            if tables_text:
                full_text += "\n\n" + "\n\n".join(tables_text)
            
            docs = [{
                self.text_column: full_text,
                "type": "text"
            }]
            
            return self.filter(docs)
        except Exception as e:
            raise ValueError(f"Failed to read Word document {file_path}: {str(e)}")

